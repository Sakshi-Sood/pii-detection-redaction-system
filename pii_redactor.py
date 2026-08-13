"""Redact PII from Red_Herring_Prospectus.docx with consistent fake replacements."""

if __name__ == "__main__":
    from docx import Document

    doc = Document("Red Herring Prospectus.docx")
    print(f"Paragraphs : {len(doc.paragraphs)}")
    print(f"Tables     : {len(doc.tables)}")
